import os

import docx
from docx import Document
from itertools import tee, islice
from docx.shared import Pt, Cm, Inches
from flask import Flask, render_template, redirect, flash, Blueprint, request, session, jsonify, current_app, url_for
from flask_login import login_required
import datetime
from .convertto import convert_to

from sqlalchemy.orm import joinedload
from werkzeug.utils import secure_filename
from cuba.extends import db
from cuba.models import *

from multiprocessing import Process

from cuba.views.commands import nnunet

import base64

medical = Blueprint("medical", __name__)


# 上下文处理器
@medical.context_processor
def inject_user_info():
    userHead = session.get('userHead')
    username = session.get('username')
    isAdmin = session.get('admin')
    return dict(username=username, isAdmin=isAdmin, userHead=userHead)


# ------------医疗界面

@medical.route('/importPicture')
@login_required
def import_picture():
    users = UserMessage.query.all()  # 获取所有用户
    ages = [user.age for user in users]  # 获取所有用户的年龄
    context = {"breadcrumb": {"parent": "3D医疗图片解析", "child": "解析医疗图片", "jsFunction": 'startTime()'}}
    return render_template("medical/importPicture/importPicture.html", **context, users=users, ages=ages)


@medical.route('/viewPicture/')
@login_required
def view_picture():
    # 页码：默认显示第一页
    page = int(request.args.get('page', 1))
    # per_page: 每页显示数据量
    per_page = int(request.args.get('per_page', 6))

    medical_pictures = (MedicalPicture.query
                        .join(User, MedicalPicture.user_id == User.id)
                        .join(UserMessage, User.id == UserMessage.user_id)
                        .options(joinedload('user'), joinedload('user.user_message'))
                        .paginate(page=page, per_page=per_page, error_out=False))
    processes = (MedicalPicture.query
                 .filter(MedicalPicture.isDoing)
                 .join(User, MedicalPicture.user_id == User.id)
                 .options(joinedload('user'), joinedload('user.user_message'))
                 .paginate(page=page, per_page=per_page, error_out=False))
    completes = (MedicalPicture.query
                 .filter(~MedicalPicture.isDoing)
                 .join(User, MedicalPicture.user_id == User.id)
                 .options(joinedload('user'), joinedload('user.user_message'))
                 .paginate(page=page, per_page=per_page, error_out=False))
    context = {"breadcrumb": {"parent": "3D医疗图片解析", "child": "查看医疗图片"}}
    return render_template("medical/viewPicture/viewPicture.html", **context, medical_pictures=medical_pictures, processes=processes, complates=completes)


@medical.route('/check_isDoing')
@login_required
def check_isDoing():
    # 查询医疗图片数据
    medical_pictures = MedicalPicture.query.join(User, (MedicalPicture.user_id == User.id)).all()
    # 构造返回给前端的数据
    medical_pictures_data = []
    for medical_picture in medical_pictures:
        medical_pictures_data.append({
            'id': medical_picture.id,
            'isDoing': medical_picture.isDoing
        })
    # 返回 JSON 格式的数据
    return jsonify({'medicalPictures': medical_pictures_data})


@medical.route('/Pictures')
@login_required
def Pictures():
    submit = request.args.get('submit')
    output = request.args.get('output')
    hide_footer = True

    # 根据 output 和 submit 查询医疗图片信息
    medical_picture = MedicalPicture.query.filter_by(outputImage=output, submitImage=submit).first()
    if medical_picture:
        user_id = medical_picture.user_id
        # 如果找到医疗图片信息，则获取用户 ID 和图片信息
        user = medical_picture.user  # 获取用户信息
        user_message = user.user_message  # 获取用户详细信息
        image_info = {
            'id': medical_picture.id,
            'imageType': medical_picture.imageType,
            'uploadTime': medical_picture.uploadTime,
            'description': medical_picture.description,
            'submitImage': medical_picture.submitImage,
            'outputImage': medical_picture.outputImage,
            'isDoing': medical_picture.isDoing,
            'user': {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'isAdmin': user.isAdmin,
                'user_message':{
                    'id': user_message.id,
                    'userHead': user_message.userHead,
                    'name': user_message.name,
                    'age': user_message.age,
                    'sex': user_message.sex,
                    'asset': user_message.asset,
                    'idCard': user_message.idCard
                }
            }
        }
    else:
        user_id = None
        image_info = None

    return render_template("medical/Pictures/Pictures.html", hide_footer=hide_footer, submit=submit, output=output, user_id=user_id, image_info=image_info)


@medical.route('/addModal', methods=['POST'])
@login_required
def add_modal():
    # 接收表单数据
    image_id = request.form.get('imageId')
    image = request.form.get('imageData')
    image_time = datetime.datetime.strptime(request.form.get('imageTime'), '%m/%d/%Y %I:%M %p')
    description = request.form.get('description')

    # 解码图片数据并保存到指定目录
    image_data = image.replace('data:image/png;base64,', '')
    image_bytes = base64.b64decode(image_data)
    formatted_time = image_time.strftime('%Y-%m-%d_%H_%M').replace(' ', '').replace(':', '_')
    image_filename = f"{formatted_time}_image.png"
    # 构建图片保存路径
    base_path = f"static/assets/images/Pictures/{image_id}/"
    save_path = os.path.join(f'/static/assets/images/Pictures/{image_id}/', image_filename)
    image_path = os.path.join(current_app.root_path, base_path, image_filename)
    # 确保目录存在，如果不存在则创建
    os.makedirs(os.path.dirname(image_path), exist_ok=True)
    # 保存图片到指定路径
    with open(image_path, 'wb') as f:
        f.write(image_bytes)

    # 保存表单数据和图片路径到数据库中
    modal_list = ModalList(
        image_id=image_id,
        image_time=image_time,
        description=description,
        image=save_path)
    db.session.add(modal_list)
    db.session.commit()

    return jsonify({'message': '记录保存成功!'}), 200


@medical.route('/deleteModal/<int:modal_id>', methods=['DELETE'])
@login_required
def delete_modal(modal_id):
    # 查询要删除的数据
    modal = ModalList.query.get(modal_id)
    if modal:
        # 删除数据
        db.session.delete(modal)
        db.session.commit()
        # 返回成功消息
        return jsonify({'message': '删除成功'}), 200
    else:
        # 如果找不到对应的数据，返回错误消息
        return jsonify({'message': '未找到要删除的数据'}), 404


@medical.route('/updateModal', methods=['POST', 'GET'])
@login_required
def updateModal():
    try:
        # 从请求表单中获取数据
        list_id = request.form.get('list_id')
        image_id = request.form.get('query_id')
        image_time = datetime.datetime.strptime(request.form.get('updateImageTime'), '%m/%d/%Y %I:%M %p')
        description = request.form.get('updateDescription')

        # 查询数据库中对应的记录
        modal = ModalList.query.get(list_id)

        # 更新记录的其他字段
        modal.image_time = image_time
        modal.description = description

        # 提交更新到数据库
        db.session.commit()

        # 返回 JSON 响应
        return jsonify({'message': '记录修改成功!'}), 200
    except Exception as e:
        # 返回 JSON 响应，表示修改失败
        return jsonify({'error': str(e)}), 500


@medical.route('/queryModal', methods=['POST', 'GET'])
@login_required
def queryModal():
    image_id = request.args.get('image_id')  # 获取前端传递的 image_id
    print(image_id)

    # 根据 image_id 查询数据库中所有数据
    modal_data = ModalList.query.filter_by(image_id=image_id).all()

    # 将查询到的数据转换成字典列表
    modal_list = []
    for modal in modal_data:
        modal_list.append({
            'id': modal.id,
            'image_id': modal.image_id,
            'image_time': modal.image_time.strftime('%m/%d/%Y %I:%M %p'),
            'description': modal.description,
            'image': modal.image
        })

    print(modal_list)
    # 返回 JSON 格式的数据
    return jsonify({'modal_list': modal_list})


@medical.route('/insertMoadlDocx', methods=['POST', 'GET'])
@login_required
def insertModalDocx():
    try:
        image_id = request.form.get('image_id')
        # 查询相关信息
        medical_picture_info = MedicalPicture.query.filter_by(id=image_id).first()
        user_info = User.query.filter_by(id=medical_picture_info.user_id).first()
        user_message_info = UserMessage.query.filter_by(user_id=user_info.id).first()
        modal_list_info = ModalList.query.filter_by(image_id=image_id).all()

        # 读取docx模板
        template_path = os.path.join(current_app.root_path, 'static', 'word', 'template.docx')
        doc = Document(template_path)

        # 替换表格占位符
        placeholders = {
            '{{username}}': user_info.username,
            '{{name}}': user_message_info.name,
            '{{sex}}': '男' if user_message_info.sex == 1 else '女',
            '{{age}}': str(user_message_info.age),
            '{{imageType}}': medical_picture_info.imageType,
            '{{uploadTime}}': str(medical_picture_info.uploadTime),
            '{{phone}}': user_message_info.phone,
            '{{idCard}}': str(user_message_info.idCard),
            '{{asset}}': user_message_info.asset
        }

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholders.items():
                        if key in cell.text:
                            # 保留原始字体格式
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

        # 循环插入ModalList信息
        for index, item in enumerate(modal_list_info):
            if index == 0:
                # 如果是第一条记录，直接替换原有的占位符
                for paragraph in doc.paragraphs:
                    if '{{description}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{description}}','\t' + item.description)
                    if '{{image}}' in paragraph.text:
                        # 删除原有的占位符
                        paragraph.text = paragraph.text.replace('{{image}}', '')
                        # 添加图片
                        run = paragraph.add_run()
                        image_path = os.path.join(current_app.root_path, item.image.lstrip('/'))
                        run.add_picture(image_path, width=docx.shared.Cm(14.5), height=docx.shared.Cm(5.2))
            else:
                # 如果不是第一条记录，在报告医师信息的上一行插入新段落并插入数据
                paragraphs_copy, paragraphs_iter = tee(doc.paragraphs)
                for i, paragraph in enumerate(paragraphs_iter):
                    if '报告医师：' in paragraph.text:
                        # 在报告医师信息的上一行插入空白行
                        doc.paragraphs[i - 1].insert_paragraph_before()

                        # 在空白行之后插入新段落
                        new_paragraph = doc.paragraphs[i].insert_paragraph_before()

                        # 插入诊断描述和图片信息
                        new_run1 = new_paragraph.add_run(f"诊断描述：\n")
                        new_run1.font.name = '宋体'  # 设置字体为宋体
                        new_run1.font.size = Pt(12)  # 设置字号为12磅
                        new_paragraph.add_run('\t')  # 添加制表符实现缩进
                        new_paragraph.add_run(item.description)
                        new_run2 =  new_paragraph.add_run(f"\n诊断图片：\n")  # 设置字体为宋体
                        new_run2.font.name = '宋体'  # 设置字体为宋体
                        new_run2.font.size = Pt(12)  # 设置字号为12磅
                        image_path = os.path.join(current_app.root_path, item.image.lstrip('/'))
                        new_paragraph.add_run().add_picture(image_path, width=docx.shared.Cm(14.5),
                                                            height=docx.shared.Cm(5.2))
                        break

                # 添加一个空行，用于分隔不同的记录
                doc.add_paragraph()

        docx_filename = f"{image_id}_{user_message_info.name}_{medical_picture_info.imageType}.docx"
        folder_name = os.path.splitext(docx_filename)[0]  # 去掉文件尾缀
        docx_folder = os.path.join(current_app.root_path, 'static', 'word', folder_name)  # 使用去掉尾缀后的文件名作为文件夹名

        # 确保文件夹存在，如果不存在则创建
        if not os.path.exists(docx_folder):
            os.makedirs(docx_folder)

        # 保存 DOCX 文件
        docx_path = os.path.join(docx_folder, docx_filename)
        doc.save(docx_path)

        # from win32com.client import pythoncom  # 导入 pythoncom
        # pythoncom.CoInitialize()  # 初始化 COM 线程
        # # 构建 PDF 文件路径
        # pdf_filename = docx_filename.replace('.docx', '.pdf')
        # pdf_folder = docx_folder  # 与 DOCX 文件相同的目录
        # pdf_path = os.path.join(pdf_folder, pdf_filename)
        #
        # # 将 DOCX 文件转换为 PDF
        # convert(docx_path, pdf_path)

        # 创建 PDF 文件
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        convert_to([docx_path], "pdf")

        # 构建目标文件的路径
        docx_save_path = os.path.join('/static', 'word', folder_name, docx_filename)
        pdf_save_path = os.path.join('/static', 'word', folder_name, 'out', pdf_filename)

        # 替换所有路径中的反斜杠为正斜杠
        docx_save_path = docx_save_path.replace('\\', '/')
        pdf_save_path = pdf_save_path.replace('\\', '/')

        # 将路径保存到数据库中
        medical_picture_info.pdf_path = pdf_save_path
        medical_picture_info.docx_path = docx_save_path

        db.session.commit()

        # 返回 JSON 响应
        return jsonify({'message': '报告生成成功!'}), 200
    except Exception as e:
        # 返回 JSON 响应，表示修改失败
        return jsonify({'error': str(e)}), 500


# -------------------File Manager
@medical.route('/file_manager/')
@login_required
def file_manager():
    # 页码：默认显示第一页
    page = int(request.args.get('page', 1))
    # per_page: 每页显示数据量
    per_page = int(request.args.get('per_page', 7))

    medical_pictures = (MedicalPicture.query
                        .join(User, MedicalPicture.user_id == User.id)
                        .join(UserMessage, User.id == UserMessage.user_id)
                        .options(joinedload('user'), joinedload('user.user_message'))
                        .paginate(page=page, per_page=per_page, error_out=False))

    context = {"breadcrumb": {"parent": "文件管理", "child": "病例报告管理"}}
    return render_template("applications/file-manager/file-manager.html", **context, medical_pictures=medical_pictures)


@medical.route('/appointment-list/')
@login_required
def appointments():
    # 页码：默认显示第一页
    page = int(request.args.get('page', 1))
    # per_page: 每页显示数据量
    per_page = int(request.args.get('per_page', 7))
    user_id = User.query.filter_by(username=session.get('username')).first().id
    doctor_id = Doctor.query.filter_by(user_id=user_id).first().id
    appointment_list = db.session.query(Appointment, Doctor, Department, UserMessage). \
        join(Doctor, Appointment.doctor_id == Doctor.id). \
        join(Department, Doctor.department_id == Department.id). \
        join(UserMessage, Appointment.user_id == UserMessage.user_id). \
        filter(Appointment.doctor_id == doctor_id). \
        paginate(page=page, per_page=per_page)

    context = {"breadcrumb": {"parent": "医疗应用", "child": "处理挂号"}}
    return render_template('medical/view-appointment/view-appointment.html', **context, appointment_list=appointment_list)


# 后台路由处理更新预约状态的请求
@medical.route('/update_appointment_status', methods=['POST'])
@login_required
def update_appointment_status():
    data = request.json
    appointment_id = data.get('appointmentId')

    # 根据预约 ID 更新预约状态为已处理
    appointment = Appointment.query.get(appointment_id)
    if appointment:
        appointment.status = 0
        db.session.commit()
        return jsonify({'success': True})
    else:
        return jsonify({'success': False}), 400


@medical.route('/medical', methods=['POST'])
@login_required
def medical_add():
    if request.method == 'POST':
        try:
            name = request.form['name']
            image_type = request.form['imageType']
            age = int(request.form['age'])
            upload_time = datetime.datetime.strptime(request.form['uploadTime'], '%Y/%m/%d %I:%M %p')
            description = request.form['description']

            # 格式化上传时间
            formatted_upload_time = upload_time.strftime('%Y-%m-%d_%H_%M').replace(' ', '').replace(':', '_')
            # 格式化文件夹名
            folder_name = f"{name}_{age}_{formatted_upload_time}"

            # 查询用户的唯一ID
            user = UserMessage.query.filter_by(name=name, age=age).first()
            if user:
                user_id = user.user_id
            else:
                return jsonify({'error': '找不到对应的用户!'}), 400

            # 处理文件上传
            if 'file' in request.files:
                files = request.files.getlist('file')  # 获取所有上传的文件列表
                if files:
                    # 构建目标文件夹路径
                    target_folder = os.path.join(current_app.root_path, 'static', 'medical', folder_name, 'submit')
                    save_folder = os.path.join('medical', folder_name, 'output')
                    # 确保目标文件夹存在
                    os.makedirs(target_folder, exist_ok=True)

                    for file in files:
                        if file.filename != '':
                            # 获取文件后缀
                            file_extension = '.' + file.filename.split('.', 1)[-1]

                            # 构建新文件名，根据文件数量递增并补齐至四位数
                            existing_files = os.listdir(target_folder)
                            file_count = len(existing_files)
                            new_filename = f"{folder_name}_{str(file_count).zfill(4)}{file_extension}"
                            save_filename = f"{folder_name}{file_extension}"

                            # 构建目标文件的路径
                            target_file_path = os.path.join(target_folder, new_filename)
                            submit_path = os.path.join('/static', 'medical', folder_name, 'submit', new_filename)
                            output_path = os.path.join('/static', save_folder, save_filename)

                            # 保存文件到目标路径
                            file.save(target_file_path)

                            # 将文件路径以及其他信息存入数据库
                            medical_picture = MedicalPicture(
                                imageType=image_type,
                                user_id=user_id,
                                uploadTime=upload_time,
                                description=description,
                                submitImage=submit_path,  # 保存目标文件的路径
                                outputImage=output_path,
                                isDoing=1
                            )
                            db.session.add(medical_picture)
                            db.session.commit()

                            # 获取刚插入到数据库的MedicalPicture的id
                            medical_picture_id = medical_picture.id

                            # 在上传文件处理完成后，启动新进程执行 nnunet() 函数
                            target_submit_folder = os.path.join(current_app.root_path, 'static', 'medical', folder_name,
                                                                'submit')
                            target_output_folder = os.path.join(current_app.root_path, 'static', 'medical', folder_name,
                                                                'output')
                            nnunet_process = Process(target=nnunet, args=(target_submit_folder, target_output_folder, user_id, medical_picture_id))
                            nnunet_process.start()

                            # 向前端发送成功响应
                            return jsonify({'message': '文件上传成功!'}), 200

            return jsonify({'error': '请求中缺少文件!'}), 400
        except Exception as e:
            print('error:', str(e))
            return jsonify({'error': str(e)}), 500  # 返回更具体的错误信息
    return jsonify({'error': '请求失败!'}), 400


